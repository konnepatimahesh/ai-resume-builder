from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def generate_docx(optimized_text, output_path):
    doc = Document()
    lines = optimized_text.split("\n")

    for raw_line in lines:
        stripped = raw_line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        # Section divider line: "---" -> add a bottom border under previous heading
        if re.fullmatch(r'-{2,}', stripped):
            _add_divider(doc)
            continue

        # Section header: short ALL-CAPS line -> heading style
        is_header = stripped.isupper() and 2 <= len(stripped) < 40 and any(c.isalpha() for c in stripped)
        if is_header:
            doc.add_heading(stripped, level=2)
            continue

        # Normal line: parse **bold** segments
        p = doc.add_paragraph()
        _add_runs_with_bold(p, stripped)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.save(output_path)
    return output_path


def _add_runs_with_bold(paragraph, line):
    line = _strip_other_markdown(line)
    segments = re.split(r'(\*\*.*?\*\*)', line)
    for seg in segments:
        if not seg:
            continue
        is_bold = seg.startswith("**") and seg.endswith("**") and len(seg) > 4
        text = seg[2:-2] if is_bold else seg
        run = paragraph.add_run(text)
        run.bold = is_bold


def _add_divider(doc):
    """Add a thin horizontal line by giving an empty paragraph a bottom border."""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_after = Pt(4)

    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0B0B0')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _strip_other_markdown(text):
    text = re.sub(r'^#{1,6}\s*', '', text)
    text = re.sub(r'(?<!\*)\*(?!\*)([^*]+?)\*(?!\*)', r'\1', text)
    return text